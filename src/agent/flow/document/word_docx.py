"""document/word_docx.py — Word .docx 与 DagPlanDocument 互转（轻量兼容）。

约定（与 Markdown 视图一致）
----------------------------
· **Heading 1** / Title：计划标题（可无 «Plan:» 前缀，导出时会补上）。
· **Heading 2** «Objective» / «Tasks»：分段标题。
· 列表段落（List Bullet / List Number 或可被识别为任务行的正文）：任务行语法与
  ``DagMarkdownIO`` 相同；若在 Word 里将任务 ID 设为粗体，会导出为 ``**id**``。

依赖：``python-docx``（写入 requirements.txt）。
"""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Inches

from agent.flow.base.types import NodeStatus

from .markdown import DagMarkdownIO, _TASK_LINE
from .model import DagPlanDocument


_CHECK_HEAD = re.compile(r"^(\[[xX\-~!> ]\]|☐|☑|✓|✔)")


class DagWordIO:
    """``.docx`` ↔ 内部 Markdown 字符串 ↔ ``DagPlanDocument``。"""

    @classmethod
    def from_docx_path(cls, path: str | Path, *, strict: bool = True) -> DagPlanDocument:
        p = Path(path).expanduser()
        dw = DocxDocument(str(p))
        md = cls._document_to_markdown(dw)
        return DagMarkdownIO.from_markdown(md, strict=strict)

    @classmethod
    def from_docx_bytes(cls, data: bytes, *, strict: bool = True) -> DagPlanDocument:
        dw = DocxDocument(BytesIO(data))
        md = cls._document_to_markdown(dw)
        return DagMarkdownIO.from_markdown(md, strict=strict)

    @classmethod
    def to_docx_path(cls, plan_doc: DagPlanDocument, path: str | Path) -> None:
        p = Path(path).expanduser()
        dw = DocxDocument()
        dw.add_heading(f"Plan: {plan_doc.title}", level=1)
        dw.add_heading("Objective", level=2)
        dw.add_paragraph(plan_doc.objective)
        dw.add_heading("Tasks", level=2)
        mark = DagMarkdownIO._STATUS_MARK[NodeStatus.pending]
        for n in plan_doc.nodes:
            ann_parts: list[str] = []
            if n.depends_on:
                ann_parts.append(f"`depends_on:{','.join(n.depends_on)}`")
            if n.tool_package:
                ann_parts.append(f"`tool:{n.tool_package}`")
            if n.max_steps is not None:
                ann_parts.append(f"`max_steps:{n.max_steps}`")
            if n.system_note:
                ann_parts.append(f"`note:{n.system_note}`")
            for k, v in sorted(n.tags.items()):
                if k.startswith("_"):
                    continue
                ann_parts.append(f"`tag:{k}={v}`")
            ann_str = (" " + " ".join(ann_parts)) if ann_parts else ""

            para = dw.add_paragraph(style="List Bullet")
            para.add_run(f"{mark} ")
            rid = para.add_run(n.task_id)
            rid.bold = True
            if ann_str:
                para.add_run(ann_str)

            desc = n.description.strip()
            if desc:
                for part in desc.splitlines():
                    body = dw.add_paragraph(part)
                    body.paragraph_format.left_indent = Inches(0.35)

            dw.add_paragraph("")

        dw.save(str(p))

    # ── OOXML → Markdown（供 DagMarkdownIO 解析） ────────────────────────────

    @staticmethod
    def _heading_level(style_name: str) -> int | None:
        if style_name == "Title":
            return 1
        if style_name.startswith("Heading"):
            parts = style_name.split()
            if len(parts) >= 2 and parts[1].isdigit():
                return int(parts[1])
        return None

    @staticmethod
    def _runs_as_inline_md(paragraph: object) -> str:
        parts: list[str] = []
        for run in paragraph.runs:
            t = run.text
            if not t:
                continue
            if run.bold:
                parts.append(f"**{t}**")
            else:
                parts.append(t)
        return "".join(parts)

    @staticmethod
    def _normalize_list_line(s: str) -> str:
        t = s.strip().replace("\u00a0", " ").replace("\u200b", "")
        if t.startswith("\uf0b7"):
            t = "- " + t[1:].lstrip()
        elif t.startswith("•"):
            t = "- " + t[1:].lstrip()
        elif t.startswith("☐"):
            t = "- [ ] " + t[1:].lstrip()
        elif t.startswith("☑"):
            t = "- [x] " + t[1:].lstrip()
        elif t.startswith("✓") or t.startswith("✔"):
            t = "- [x] " + t[1:].lstrip()
        elif t.startswith("□"):
            rest = t[1:].lstrip()
            t = "- [ ] " + rest
        if _CHECK_HEAD.match(t) and not t.lstrip().startswith("-"):
            t = "- " + t.lstrip()
        return t

    @classmethod
    def _document_to_markdown(cls, dw: DocxDocument) -> str:
        lines: list[str] = []
        for p in dw.paragraphs:
            style_name = p.style.name if p.style else ""
            hl = cls._heading_level(style_name)
            plain = p.text.strip()

            if hl == 1:
                t = plain
                if not t.lower().startswith("plan:"):
                    t = "Plan: " + t
                lines.append("# " + t)
                continue
            if hl == 2:
                lines.append(f"## {plain}")
                continue
            if hl is not None and hl >= 3:
                lines.append("#" * hl + " " + plain)
                continue

            inline = cls._runs_as_inline_md(p).rstrip()
            if not inline.strip():
                lines.append("")
                continue

            stl = style_name.lower()
            if "list" in stl:
                lines.append(cls._normalize_list_line(inline))
                continue

            cand = cls._normalize_list_line(inline)
            if _TASK_LINE.match(cand.strip()):
                lines.append(cand.strip())
                continue

            lines.append(inline)

        return "\n".join(lines).rstrip() + "\n"
